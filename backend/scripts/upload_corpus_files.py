#!/usr/bin/env python3
"""
Upload Corpus Files lên MinIO + PostgreSQL + Redis

Script này cho phép upload các file .txt, .docx, .pdf từ thư mục cục bộ trực tiếp vào hệ thống.

Cách sử dụng:
    # Upload tất cả file từ một thư mục
    python scripts/upload_corpus_files.py --dir /path/to/corpus/files

    # Upload kèm metadata tùy chỉnh
    python scripts/upload_corpus_files.py --dir /path/to/files --author "Nguyễn Văn A" --university "ĐH CNTT"

    # Upload và đồng bộ ngay vào Redis
    python scripts/upload_corpus_files.py --dir /path/to/files --sync-redis

    # Giới hạn số lượng file
    python scripts/upload_corpus_files.py --dir /path/to/files --limit 100

Ví dụ:
    # Test nhanh: upload 100 file txt
    python scripts/upload_corpus_files.py --dir /app/corpus_data --limit 100 --sync-redis
"""
import os
import sys
import uuid
import hashlib
import argparse
from datetime import datetime
import logging

sys.path.insert(0, os.path.dirname(os.path.dirname(os.path.abspath(__file__))))

from app.config import settings
from app.db.database import SessionLocal
from app.db.models import Document
from app.services.minio_storage import get_minio_storage
from app.services.preprocessing.vietnamese_nlp import preprocess_vietnamese
from app.services.preprocessing.text_normalizer import normalize_text

logging.basicConfig(level=logging.INFO, format='%(asctime)s - %(levelname)s - %(message)s')
logger = logging.getLogger(__name__)

SUPPORTED_EXTENSIONS = {'.txt', '.docx', '.pdf', '.doc'}


def extract_text_from_file(filepath: str) -> str:
    """Trích xuất nội dung văn bản từ file"""
    ext = os.path.splitext(filepath)[1].lower()
    
    if ext == '.txt':
        with open(filepath, 'r', encoding='utf-8', errors='ignore') as f:
            return f.read()
    
    elif ext == '.docx':
        try:
            import docx
            doc = docx.Document(filepath)
            return '\n'.join(p.text for p in doc.paragraphs if p.text.strip())
        except ImportError:
            logger.warning("python-docx chưa được cài đặt. Chạy lệnh: pip install python-docx")
            return ""
    
    elif ext == '.pdf':
        try:
            import PyPDF2
            with open(filepath, 'rb') as f:
                reader = PyPDF2.PdfReader(f)
                return '\n'.join(page.extract_text() or '' for page in reader.pages)
        except ImportError:
            logger.warning("PyPDF2 chưa được cài đặt. Chạy lệnh: pip install PyPDF2")
            return ""
    
    return ""


def upload_corpus(args):
    """Upload các file corpus từ thư mục vào hệ thống"""
    directory = args.dir
    if not os.path.isdir(directory):
        logger.error(f"❌ Không tìm thấy thư mục: {directory}")
        sys.exit(1)
    
    # Thu thập các file hợp lệ
    files = []
    for f in sorted(os.listdir(directory)):
        ext = os.path.splitext(f)[1].lower()
        if ext in SUPPORTED_EXTENSIONS:
            files.append(os.path.join(directory, f))
    
    if args.limit:
        files = files[:args.limit]
    
    if not files:
        logger.error(f"❌ Không tìm thấy file hỗ trợ nào trong thư mục {directory}")
        sys.exit(1)
    
    logger.info(f"\n{'='*70}")
    logger.info(f"📤 ĐANG UPLOAD {len(files)} FILE CORPUS")
    logger.info(f"   Nguồn: {directory}")
    logger.info(f"{'='*70}\n")
    
    db = SessionLocal()
    minio = get_minio_storage()
    minio_available = minio.is_available()
    
    if not minio_available:
        logger.warning("⚠️  MinIO chưa khả dụng. File sẽ chỉ được lưu vào PostgreSQL.")
    
    uploaded = 0
    skipped = 0
    seen_hashes = set()
    
    for i, filepath in enumerate(files, 1):
        filename = os.path.basename(filepath)
        try:
            # Trích xuất văn bản
            text = extract_text_from_file(filepath)
            if not text or len(text.strip()) < 50:
                logger.warning(f"⚠️  [{i}/{len(files)}] {filename} - Quá ngắn hoặc rỗng, bỏ qua")
                skipped += 1
                continue
            
            # Kiểm tra trùng lặp
            text_hash = hashlib.sha256(text.encode()).hexdigest()
            if text_hash in seen_hashes:
                logger.warning(f"⚠️  [{i}/{len(files)}] {filename} - Trùng lặp trong batch, bỏ qua")
                skipped += 1
                continue
            
            existing = db.query(Document).filter(Document.file_hash_sha256 == text_hash).first()
            if existing:
                logger.warning(f"⚠️  [{i}/{len(files)}] {filename} - Đã tồn tại trong DB, bỏ qua")
                seen_hashes.add(text_hash)
                skipped += 1
                continue
            
            # Xử lý văn bản
            normalized = normalize_text(text)
            tokens = preprocess_vietnamese(normalized)
            word_count = len(tokens)
            
            # Tạo tiêu đề từ tên file
            title = os.path.splitext(filename)[0].replace('_', ' ').replace('-', ' ')
            doc_id = uuid.uuid4()
            year = args.year or datetime.now().year
            
            # Lưu vào PostgreSQL
            doc = Document(
                id=doc_id,
                title=title[:500],
                author=args.author or 'Unknown',
                university=args.university or 'Unknown',
                year=year,
                original_filename=filename,
                file_hash_sha256=text_hash,
                file_size_bytes=os.path.getsize(filepath),
                word_count=word_count,
                language='vi',
                extraction_method='file_upload',
                extracted_text=text,
                is_corpus=1,
                status='indexed',
                indexed_at=datetime.now()
            )
            
            db.add(doc)
            try:
                db.commit()
            except Exception as e:
                db.rollback()
                logger.warning(f"⚠️  [{i}/{len(files)}] {filename} - Lỗi database: {str(e)[:60]}")
                skipped += 1
                continue
            
            # Upload lên MinIO
            if minio_available:
                minio.upload_corpus_document(
                    doc_id=str(doc_id),
                    title=title,
                    text=text,
                    author=args.author or 'Unknown',
                    university=args.university or 'Unknown',
                    year=year
                )
            
            seen_hashes.add(text_hash)
            uploaded += 1
            logger.info(f"✅ [{uploaded}/{len(files)}] {filename} - {word_count} từ")
        
        except Exception as e:
            logger.error(f"❌ [{i}/{len(files)}] {filename} - Lỗi: {e}")
            try:
                db.rollback()
            except Exception:
                pass
            skipped += 1
    
    db.close()
    
    logger.info(f"\n{'='*70}")
    logger.info(f"✅ Tóm tắt upload:")
    logger.info(f"   • Đã upload: {uploaded} file")
    logger.info(f"   • Bỏ qua: {skipped} file")
    logger.info(f"   • MinIO: {'✅' if minio_available else '❌ Không khả dụng'}")
    logger.info(f"{'='*70}\n")
    
    # Đồng bộ Redis
    if args.sync_redis and uploaded > 0:
        logger.info("🔄 Đang restart backend để đồng bộ Redis LSH index...")
        import subprocess
        try:
            subprocess.run(['docker', 'restart', 'plagiarism-backend'], check=True, timeout=30)
            logger.info("✅ Backend đã restart. Redis sẽ đồng bộ khi khởi động lại.")
        except Exception:
            logger.info("⚠️  Không thể restart tự động. Hãy chạy thủ công: docker restart plagiarism-backend")


def main():
    parser = argparse.ArgumentParser(description='Upload file corpus lên MinIO + PostgreSQL')
    parser.add_argument('--dir', required=True, help='Thư mục chứa các file corpus (.txt, .docx, .pdf)')
    parser.add_argument('--limit', type=int, help='Giới hạn số lượng file upload')
    parser.add_argument('--author', type=str, help='Tên tác giả áp dụng cho tất cả file')
    parser.add_argument('--university', type=str, help='Tên trường đại học áp dụng cho tất cả file')
    parser.add_argument('--year', type=int, help='Năm xuất bản')
    parser.add_argument('--sync-redis', action='store_true', help='Đồng bộ vào Redis sau khi upload')
    args = parser.parse_args()
    upload_corpus(args)


if __name__ == '__main__':
    main()